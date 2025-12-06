"""Document processing service.

Supports:
- PDF text extraction
- DOCX text extraction
- Plain text files
- Basic OCR for scanned PDFs (optional)
"""

import base64
import logging
from dataclasses import dataclass
from io import BytesIO
from typing import Any

logger = logging.getLogger(__name__)


@dataclass
class DocumentInfo:
    """Information about a document."""

    format: str
    mime_type: str
    page_count: int | None = None
    word_count: int = 0
    size_bytes: int = 0
    title: str | None = None
    author: str | None = None


@dataclass
class ExtractedDocument:
    """Extracted content from a document."""

    text: str
    info: DocumentInfo
    pages: list[str] | None = None  # Per-page text for PDFs
    metadata: dict[str, Any] | None = None


class DocumentProcessingError(Exception):
    """Raised when document processing fails."""

    pass


class DocumentProcessor:
    """
    Processes documents for text extraction.

    Features:
    - PDF text extraction (PyPDF2/pypdf)
    - DOCX text extraction (python-docx)
    - Plain text file handling
    - Word/page count statistics
    """

    MIME_TYPES = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "doc": "application/msword",
        "txt": "text/plain",
        "md": "text/markdown",
        "rtf": "application/rtf",
    }

    FORMAT_SIGNATURES = {
        b"%PDF": "pdf",
        b"PK\x03\x04": "docx",  # Also could be xlsx, pptx - check further
    }

    def __init__(
        self,
        max_size_bytes: int = 20 * 1024 * 1024,  # 20MB
        max_pages: int = 100,
    ):
        self._max_size = max_size_bytes
        self._max_pages = max_pages

    def _detect_format(self, data: bytes) -> str | None:
        """Detect document format from magic bytes."""
        if len(data) < 8:
            return None

        # Check PDF
        if data[:4] == b"%PDF":
            return "pdf"

        # Check ZIP-based formats (DOCX, XLSX, PPTX)
        if data[:4] == b"PK\x03\x04":
            # Need to check internal structure
            try:
                import zipfile
                with zipfile.ZipFile(BytesIO(data)) as zf:
                    names = zf.namelist()
                    if "word/document.xml" in names:
                        return "docx"
                    if "xl/workbook.xml" in names:
                        return "xlsx"
                    if "ppt/presentation.xml" in names:
                        return "pptx"
            except Exception:
                pass
            return "docx"  # Default to docx for Office Open XML

        # Check RTF
        if data[:5] == b"{\\rtf":
            return "rtf"

        # Check if it might be plain text
        try:
            sample = data[:1000].decode("utf-8")
            if sample.isprintable() or "\n" in sample:
                return "txt"
        except UnicodeDecodeError:
            pass

        return None

    def get_info(self, doc_bytes: bytes) -> DocumentInfo:
        """
        Get information about a document.

        Args:
            doc_bytes: Raw document bytes

        Returns:
            DocumentInfo with format, page count, etc.
        """
        if len(doc_bytes) > self._max_size:
            raise DocumentProcessingError(
                f"Document size {len(doc_bytes) / 1024 / 1024:.1f}MB "
                f"exceeds maximum {self._max_size / 1024 / 1024:.0f}MB"
            )

        format = self._detect_format(doc_bytes)
        if not format:
            raise DocumentProcessingError("Unsupported or unrecognized document format")

        mime_type = self.MIME_TYPES.get(format, f"application/{format}")

        info = DocumentInfo(
            format=format,
            mime_type=mime_type,
            size_bytes=len(doc_bytes),
        )

        # Get additional info based on format
        if format == "pdf":
            self._get_pdf_info(doc_bytes, info)
        elif format == "docx":
            self._get_docx_info(doc_bytes, info)

        return info

    def _get_pdf_info(self, data: bytes, info: DocumentInfo) -> None:
        """Extract PDF metadata."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(data))
            info.page_count = len(reader.pages)

            if reader.metadata:
                info.title = reader.metadata.title
                info.author = reader.metadata.author

        except ImportError:
            logger.warning("pypdf not installed, limited PDF support")
        except Exception as e:
            logger.debug(f"PDF info extraction failed: {e}")

    def _get_docx_info(self, data: bytes, info: DocumentInfo) -> None:
        """Extract DOCX metadata."""
        try:
            from docx import Document

            doc = Document(BytesIO(data))

            # Count pages (approximate from paragraphs)
            paragraph_count = len(doc.paragraphs)
            info.page_count = max(1, paragraph_count // 30)  # Rough estimate

            # Get core properties
            if doc.core_properties:
                info.title = doc.core_properties.title
                info.author = doc.core_properties.author

        except ImportError:
            logger.warning("python-docx not installed, limited DOCX support")
        except Exception as e:
            logger.debug(f"DOCX info extraction failed: {e}")

    def extract_text(self, doc_bytes: bytes) -> ExtractedDocument:
        """
        Extract text from a document.

        Args:
            doc_bytes: Raw document bytes

        Returns:
            ExtractedDocument with text and metadata
        """
        info = self.get_info(doc_bytes)

        if info.format == "pdf":
            text, pages = self._extract_pdf_text(doc_bytes)
        elif info.format == "docx":
            text, pages = self._extract_docx_text(doc_bytes)
        elif info.format in ("txt", "md"):
            text = doc_bytes.decode("utf-8", errors="replace")
            pages = None
        else:
            raise DocumentProcessingError(
                f"Text extraction not supported for format: {info.format}"
            )

        # Calculate word count
        info.word_count = len(text.split())

        return ExtractedDocument(
            text=text,
            info=info,
            pages=pages,
        )

    def _extract_pdf_text(self, data: bytes) -> tuple[str, list[str] | None]:
        """Extract text from PDF."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(data))
            pages: list[str] = []
            all_text: list[str] = []

            for i, page in enumerate(reader.pages[:self._max_pages]):
                page_text = page.extract_text() or ""
                pages.append(page_text)
                all_text.append(page_text)

            return "\n\n".join(all_text), pages

        except ImportError:
            raise DocumentProcessingError(
                "pypdf not installed. Install with: pip install pypdf"
            ) from None
        except Exception as e:
            raise DocumentProcessingError(f"PDF text extraction failed: {e}") from e

    def _extract_docx_text(self, data: bytes) -> tuple[str, list[str] | None]:
        """Extract text from DOCX."""
        try:
            from docx import Document

            doc = Document(BytesIO(data))
            paragraphs: list[str] = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            return "\n\n".join(paragraphs), None

        except ImportError:
            raise DocumentProcessingError(
                "python-docx not installed. Install with: pip install python-docx"
            ) from None
        except Exception as e:
            raise DocumentProcessingError(f"DOCX text extraction failed: {e}") from e

    def extract_summary(
        self,
        doc_bytes: bytes,
        max_chars: int = 2000,
    ) -> str:
        """
        Extract a summary (first portion) of document text.

        Args:
            doc_bytes: Raw document bytes
            max_chars: Maximum characters to return

        Returns:
            Summary text
        """
        result = self.extract_text(doc_bytes)
        text = result.text.strip()

        if len(text) <= max_chars:
            return text

        # Find a good break point
        truncated = text[:max_chars]

        # Try to break at sentence or paragraph
        for sep in ["\n\n", "\n", ". ", " "]:
            last_sep = truncated.rfind(sep)
            if last_sep > max_chars * 0.7:
                return truncated[:last_sep + len(sep)].strip() + "..."

        return truncated.strip() + "..."

    def to_base64(self, doc_bytes: bytes) -> str:
        """Convert document bytes to base64 string."""
        return base64.b64encode(doc_bytes).decode("utf-8")

    def from_base64(self, base64_data: str) -> bytes:
        """Decode base64 document data."""
        if "," in base64_data:
            base64_data = base64_data.split(",", 1)[1]

        try:
            return base64.b64decode(base64_data)
        except Exception as e:
            raise DocumentProcessingError(f"Invalid base64 encoding: {e}") from e


# Global instance
_document_processor: DocumentProcessor | None = None


def get_document_processor() -> DocumentProcessor:
    """Get or create the global document processor."""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor
